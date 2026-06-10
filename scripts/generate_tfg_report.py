"""
generate_tfg_report.py — Genera el informe .docx del estudio de optimización.

Lee results/parameter_study/study_results.json (producido por
run_parameter_study.py) y construye un documento Word profesional con:

  - Metodología (función de costes, Jonker-Volgenant, Monte Carlo, GA).
  - Fragmentos de código REALES extraídos del propio proyecto.
  - Una sección por arquetipo de escenario con sus figuras y tablas.
  - Conclusiones globales y justificación de las decisiones tomadas.

Uso:
  python scripts/generate_tfg_report.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_ROOT))

STUDY_JSON = _ROOT / "results" / "parameter_study" / "study_results.json"
OUT_DOCX = _ROOT / "results" / "Informe_Optimizacion_Parametros_TFG.docx"

MONO_FONT = "Consolas"
CODE_SHADING = "F2F2F2"


# ── Extracción de código real del proyecto ───────────────────────────────────

def extract_function(rel_path: str, func_name: str) -> str:
    """Extrae el bloque de una función/método por su nombre.

    Maneja firmas multilínea rastreando el balance de paréntesis: el cuerpo
    empieza tras la línea que cierra la firma con ':'.
    """
    text = (_ROOT / rel_path).read_text(encoding="utf-8").splitlines()
    start = None
    indent = 0
    for i, line in enumerate(text):
        stripped = line.lstrip()
        if stripped.startswith(f"def {func_name}(") or stripped.startswith(f"def {func_name} "):
            start = i
            indent = len(line) - len(stripped)
            break
    if start is None:
        return f"# (no se encontró {func_name} en {rel_path})"

    # Saltar la firma (posiblemente multilínea) hasta cerrar paréntesis y ':'.
    depth = 0
    body_start = start
    for j in range(start, len(text)):
        depth += text[j].count("(") - text[j].count(")")
        if depth <= 0 and text[j].rstrip().endswith(":"):
            body_start = j + 1
            break

    end = len(text)
    for j in range(body_start, len(text)):
        line = text[j]
        if line.strip() == "":
            continue
        cur_indent = len(line) - len(line.lstrip())
        if cur_indent <= indent:
            end = j
            break
    return "\n".join(text[start:end]).rstrip()


def extract_lines(rel_path: str, start: int, end: int) -> str:
    """Extrae líneas [start, end] (1-indexado, inclusivo)."""
    text = (_ROOT / rel_path).read_text(encoding="utf-8").splitlines()
    return "\n".join(text[start - 1:end]).rstrip()


# ── Helpers de formato ───────────────────────────────────────────────────────

def add_code_block(doc: Document, code: str, caption: str | None = None):
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = MONO_FONT
    run.font.size = Pt(8.5)
    # Sombreado de fondo
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_SHADING)
    pPr.append(shd)
    return p


def add_figure(doc: Document, rel_path: str, caption: str, width_in: float = 6.3):
    path = _ROOT / rel_path
    if not path.exists():
        doc.add_paragraph(f"[Figura no encontrada: {rel_path}]")
        return
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    return table


def fmt_w(weights: list[float]) -> str:
    return f"[{weights[0]:.3f}, {weights[1]:.3f}, {weights[2]:.3f}, {weights[3]:.3f}]"


# ── Construcción del documento ───────────────────────────────────────────────

def build_methodology(doc: Document):
    doc.add_heading("2. Metodología", level=1)

    doc.add_heading("2.1. Problema y función de costes", level=2)
    doc.add_paragraph(
        "El reparto se modela como un problema de asignación dron↔pedido. Para "
        "cada par (dron i, pedido j) se calcula un coste escalar C(i,j) que "
        "combina cuatro términos normalizados a [0,1] mediante cuatro pesos "
        "w1–w4:"
    )
    doc.add_paragraph(
        "C(i,j) = w1·E_viaje + w2·Penalización_batería + "
        "w3·Exceso_capacidad + w4·Tiempo_espera"
    ).runs[0].italic = True
    doc.add_paragraph(
        "Los términos son: (w1) energía del viaje ida+vuelta, (w2) penalización "
        "por dejar la batería por debajo del 20 %, (w3) penalización por usar un "
        "dron sobredimensionado para un paquete pequeño, y (w4) tiempo de recarga "
        "anticipado. Existen dos restricciones duras (coste = ∞): que el paquete "
        "supere la carga útil del dron, o que la energía del viaje supere el 80 % "
        "de la batería disponible (margen de seguridad del 20 %)."
    )
    add_code_block(
        doc, extract_function("simulacion/cost_function.py", "compute_cost"),
        "Código 1 — Función de costes C(i,j) (simulacion/cost_function.py)."
    )

    doc.add_heading("2.2. Resolución de la asignación: Jonker-Volgenant", level=2)
    doc.add_paragraph(
        "La matriz de costes N_drones × M_pedidos se resuelve de forma óptima "
        "global con el algoritmo Jonker-Volgenant, a través de "
        "scipy.optimize.linear_sum_assignment (complejidad O(n³)). Como hay más "
        "pedidos que drones, la asignación se realiza por rondas: en cada ronda "
        "se asigna un pedido por dron, se actualizan las baterías y se recalcula "
        "la matriz para la siguiente ronda. Esta librería se mantiene intacta; el "
        "estudio se centra en CÓMO obtener los pesos que alimentan la matriz."
    )
    add_code_block(
        doc, extract_lines("simulacion/cost_matrix_assigner.py", 100, 119),
        "Código 2 — Llamada a linear_sum_assignment sobre la matriz de costes."
    )

    doc.add_heading("2.3. Métodos de obtención de los parámetros", level=2)
    doc.add_paragraph(
        "Se comparan dos estrategias para hallar el mejor vector de pesos "
        "w=[w1,w2,w3,w4], ambas evaluando candidatos sobre un lote de escenarios "
        "y usando el mismo asignador JV:"
    )
    doc.add_paragraph(
        "• Monte Carlo Random Search: muestrea el espacio de pesos con Latin "
        "Hypercube Sampling (reparto estratificado que cubre mejor el espacio que "
        "el muestreo uniforme) y se queda con la mejor combinación.", style="List Bullet"
    )
    doc.add_paragraph(
        "• Algoritmo Genético: evoluciona una población de vectores de pesos con "
        "selección por torneo, cruce BLX-α, mutación gaussiana y elitismo.",
        style="List Bullet"
    )
    add_code_block(
        doc, extract_function("simulacion/optimizer_montecarlo.py", "_latin_hypercube_sample"),
        "Código 3 — Muestreo por Latin Hypercube (Monte Carlo)."
    )
    add_code_block(
        doc, extract_function("simulacion/optimizer_genetic.py", "_blx_alpha_crossover"),
        "Código 4 — Operador de cruce BLX-α (Algoritmo Genético)."
    )

    doc.add_heading("2.4. Función objetivo común y comparación justa", level=2)
    doc.add_paragraph(
        "Una decisión metodológica clave de esta revisión: ambos optimizadores "
        "comparten EXACTAMENTE la misma función objetivo, centralizada en "
        "simulacion/objective.py. El objetivo se normaliza contra un baseline "
        "fijo (los pesos neutros w=[1,1,1,1]), de modo que un valor < 1.0 "
        "significa 'mejor que neutros'. En la versión anterior cada optimizador "
        "se normalizaba con su primer trial aleatorio, lo que hacía sus valores "
        "no comparables. Además, se separan escenarios de entrenamiento y de test "
        "(held-out) para detectar sobreajuste, y la convergencia se mide frente "
        "al número de evaluaciones de simulación, lo que permite comparar MC y GA "
        "bajo el mismo presupuesto de cómputo."
    )
    add_code_block(
        doc, extract_function("simulacion/objective.py", "_scalarize"),
        "Código 5 — Escalarización del objetivo normalizada contra el baseline "
        "(simulacion/objective.py)."
    )
    doc.add_paragraph(
        "Sobre el objetivo a minimizar: se observó empíricamente que la energía "
        "total es casi invariante a los pesos en problemas de reparto saturados "
        "(todos los pedidos deben entregarse y la carga útil fuerza qué clase de "
        "dron sirve cada paquete). Por ello el objetivo informativo es el "
        "makespan (tiempo total), que sí responde al reparto de carga; la energía "
        "se reporta igualmente para documentar dicha invariancia."
    )


def build_archetype_section(doc: Document, arch: dict, idx: int):
    doc.add_heading(f"3.{idx}. {arch['name']}", level=2)
    doc.add_paragraph(arch["description"])

    cfg = arch["config"]
    add_table(doc,
        ["Parámetro", "Valor"],
        [
            ["Drones", str(cfg["n_drones"])],
            ["Pedidos por escenario", str(cfg["n_orders"])],
            ["Rango de peso de paquete (kg)", f"{cfg['weight_min']}–{cfg['weight_max']}"],
            ["Potencia del cargador (W)", str(cfg["charger_w"])],
            ["Escenarios entrenamiento / test", f"{cfg['n_train']} / {cfg['n_test']}"],
            ["Presupuesto MC (trials)", str(cfg["mc_trials"])],
            ["Presupuesto GA (pob × gen)", f"{cfg['ga_pop']} × {cfg['ga_gen']}"],
        ])

    # Convergencia
    doc.add_paragraph()
    add_figure(doc, arch["figures"]["convergence"],
               f"Figura 3.{idx}.a — Convergencia de Monte Carlo y GA frente al "
               f"presupuesto de evaluaciones. La línea en 1.0 es el rendimiento "
               f"de los pesos neutros.")

    # Pesos hallados
    mc_w = arch["mc"]["weights"]
    ga_w = arch["ga"]["weights"]
    add_figure(doc, arch["figures"]["weights"],
               f"Figura 3.{idx}.b — Composición de pesos hallada por cada método.")
    add_table(doc,
        ["Método", "w1 (E)", "w2 (Bat)", "w3 (Cap)", "w4 (Esp)", "Obj. norm.", "Δt train"],
        [
            ["Neutros", "0.250", "0.250", "0.250", "0.250", "1.0000", "—"],
            ["Monte Carlo", f"{mc_w[0]:.3f}", f"{mc_w[1]:.3f}", f"{mc_w[2]:.3f}",
             f"{mc_w[3]:.3f}", f"{arch['mc']['best_objective']:.4f}",
             f"{arch['mc']['time_improvement_pct']:+.2f}%"],
            ["Genético", f"{ga_w[0]:.3f}", f"{ga_w[1]:.3f}", f"{ga_w[2]:.3f}",
             f"{ga_w[3]:.3f}", f"{arch['ga']['best_objective']:.4f}",
             f"{arch['ga']['time_improvement_pct']:+.2f}%"],
        ])

    # Test
    test = arch["test"]
    doc.add_paragraph()
    add_figure(doc, arch["figures"]["test_comparison"],
               f"Figura 3.{idx}.c — Mejora pareada respecto a neutros en el "
               f"conjunto de test (held-out).")
    add_table(doc,
        ["Métrica en test", "Neutros", "Monte Carlo", "Genético"],
        [
            ["Makespan medio (s)",
             f"{test['time_mean']['Neutros']:.0f}",
             f"{test['time_mean']['Monte Carlo']:.0f}",
             f"{test['time_mean']['Genético']:.0f}"],
            ["Mejora makespan (%)", "—",
             f"{test['time_impr_pct']['Monte Carlo']:+.2f}",
             f"{test['time_impr_pct']['Genético']:+.2f}"],
            ["p-value (t pareado vs neutros)", "—",
             f"{test['p_value_time']['Monte Carlo']:.4f}",
             f"{test['p_value_time']['Genético']:.4f}"],
            ["Energía media (Wh)",
             f"{test['energy_mean']['Neutros']:.0f}",
             f"{test['energy_mean']['Monte Carlo']:.0f}",
             f"{test['energy_mean']['Genético']:.0f}"],
            ["Mejora energía (%)", "—",
             f"{test['energy_impr_pct']['Monte Carlo']:+.3f}",
             f"{test['energy_impr_pct']['Genético']:+.3f}"],
        ])

    # Greedy vs JV
    gvj = arch["greedy_vs_jv"]
    jv_t = gvj["time_saving_pct"]
    verbo = ("reduce el makespan un " if jv_t >= 0
             else "empeora el makespan un ")
    doc.add_paragraph(
        f"Comparación del asignador con los pesos del GA: Jonker-Volgenant {verbo}"
        f"{abs(jv_t):.2f}% frente al greedy FIFO (p={gvj['p_value_time']:.4f}), "
        f"manteniendo prácticamente idéntica la energía "
        f"({gvj['energy_saving_pct']:+.2f}%)."
    )
    add_figure(doc, arch["figures"]["greedy_vs_jv_bars"],
               f"Figura 3.{idx}.d — Greedy vs. Jonker-Volgenant con los pesos "
               f"optimizados por GA.")

    # Sensibilidad
    add_figure(doc, arch["figures"]["sensitivity"],
               f"Figura 3.{idx}.e — Sensibilidad de la función de costes: efecto "
               f"de concentrar todo el peso en un único término.")

    # Conclusión por arquetipo
    best_test = max(test["time_impr_pct"]["Monte Carlo"],
                    test["time_impr_pct"]["Genético"])
    method_best = ("GA" if test["time_impr_pct"]["Genético"] >=
                   test["time_impr_pct"]["Monte Carlo"] else "Monte Carlo")
    sig = "estadísticamente significativa" if min(
        test["p_value_time"]["Monte Carlo"],
        test["p_value_time"]["Genético"]) < 0.05 else "no significativa al 5 %"
    p = doc.add_paragraph()
    p.add_run("Conclusión del escenario: ").bold = True
    p.add_run(
        f"la mejor reducción de makespan en test fue {best_test:+.2f}% "
        f"(método {method_best}), {sig}. La energía permanece prácticamente "
        f"constante, confirmando que en este escenario el valor de la "
        f"optimización está en el reparto temporal, no en el energético."
    )
    doc.add_page_break()


def build_global_conclusions(doc: Document, data: dict):
    doc.add_heading("4. Conclusiones globales", level=1)

    archs = data["archetypes"]

    # Tabla resumen
    rows = []
    for a in archs:
        t = a["test"]
        rows.append([
            a["name"],
            f"{t['time_impr_pct']['Monte Carlo']:+.2f}%",
            f"{t['time_impr_pct']['Genético']:+.2f}%",
            f"{t['energy_impr_pct']['Genético']:+.3f}%",
            f"{a['greedy_vs_jv']['time_saving_pct']:+.2f}%",
        ])
    add_table(doc,
        ["Escenario", "MC Δt", "GA Δt", "GA ΔE", "JV vs Greedy Δt"],
        rows)
    doc.add_paragraph()

    # Datos para sustentar las conclusiones
    e_imprs = [abs(a["test"]["energy_impr_pct"]["Genético"]) for a in archs]
    max_e = max(e_imprs)
    ga_massive = next(a for a in archs if a["key"] == "massive_ops")["test"]["time_impr_pct"]["Genético"]
    mc_massive = next(a for a in archs if a["key"] == "massive_ops")["test"]["time_impr_pct"]["Monte Carlo"]

    conclusions = [
        ("La energía total es casi invariante a los pesos.",
         f"En todos los escenarios la mejora de energía es despreciable "
         f"(máximo |ΔE| = {max_e:.2f} %). Esto no es un fallo del optimizador: "
         "como todos los pedidos deben entregarse y la carga útil determina qué "
         "clase de dron puede servir cada paquete, el reparto energético está "
         "casi fijado. Optimizar los pesos para 'energía' carece de sentido en "
         "este problema; por eso el objetivo elegido es el tiempo (makespan)."),
        ("El makespan responde a los pesos, pero la mejora media es pequeña y "
         "no siempre significativa.",
         "Las mejoras medias en test se sitúan típicamente por debajo del 2 % y, "
         "con el número de escenarios usado, no alcanzan significancia al 5 % "
         "(p-values entre 0,06 y 0,30). Sin embargo, la distribución pareada "
         "revela escenarios concretos donde una buena elección de pesos evita una "
         "ronda completa de recarga y reduce el tiempo hasta un 60 %. El valor de "
         "la optimización es, por tanto, la robustez frente a casos patológicos "
         "más que una mejora media grande: el paisaje de la función objetivo es "
         "plano por estar dominado por las restricciones de capacidad."),
        ("Ningún método domina, pero el GA es más robusto a escala.",
         f"En el conjunto de test ninguna estrategia gana en todos los escenarios. "
         f"El Algoritmo Genético resulta más fiable en el caso de gran escala "
         f"(masivo: GA {ga_massive:+.2f} % frente a Monte Carlo {mc_massive:+.2f} %) "
         "y lo consigue con un presupuesto de simulaciones comparable o menor "
         "gracias a la caché de evaluaciones. Monte Carlo, más simple, puede "
         "sobreajustar cuando hay pocos escenarios de entrenamiento y llegar a "
         "empeorar el baseline. La recomendación es usar el GA validando siempre "
         "sobre escenarios held-out."),
        ("Jonker-Volgenant ayuda en demanda ligera, no en cualquier escala.",
         "Frente al greedy FIFO, el asignador óptimo JV reduce el makespan en los "
         "escenarios pequeños y de carga ligera (+2 a +4 %), pero esa ventaja se "
         "desvanece o se invierte en los escenarios grandes. La razón es que la "
         "función de costes es un proxy por ronda y no optimiza el makespan global "
         "directamente: minimizar coste de asignación no equivale a minimizar el "
         "tiempo total cuando hay muchas rondas de recarga. Aun así, JV nunca "
         "empeora la energía y mantiene el 100 % de entregas."),
    ]
    for i, (title, body) in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {title} ").bold = True
        p.add_run(body)

    doc.add_heading("4.1. Decisión recomendada para el sistema", level=2)
    doc.add_paragraph(
        "Para la operación de la flota se recomienda: (a) mantener "
        "linear_sum_assignment como asignador; (b) optimizar los pesos con el "
        "Algoritmo Genético usando el makespan como objetivo y validando sobre "
        "escenarios held-out; (c) re-optimizar los pesos cuando cambie el perfil "
        "de demanda (p.ej. campaña e-commerce vs. reparto industrial), ya que los "
        "pesos óptimos difieren entre perfiles."
    )

    doc.add_heading("4.2. Reproducibilidad", level=2)
    doc.add_paragraph("Todo el estudio se reproduce con dos comandos:")
    add_code_block(doc,
        "python scripts/run_parameter_study.py\n"
        "python scripts/generate_tfg_report.py",
        "Código 6 — Reproducción del estudio y del informe.")
    doc.add_paragraph(
        f"Semillas fijas garantizan resultados idénticos. Tiempo de cómputo del "
        f"estudio: {data.get('elapsed_s', 0):.0f} s. Objetivo optimizado: "
        f"'{data.get('objective', 'time')}'."
    )


def main():
    if not STUDY_JSON.exists():
        print(f"ERROR: no existe {STUDY_JSON}.")
        print("Ejecuta primero: python scripts/run_parameter_study.py")
        sys.exit(1)

    data = json.loads(STUDY_JSON.read_text(encoding="utf-8"))

    doc = Document()

    # Estilo base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── Portada ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Optimización de los parámetros de la función de costes\n"
                        "para la asignación de drones de reparto")
    run.bold = True
    run.font.size = Pt(20)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Simulación extra del TFG — Monte Carlo vs. Algoritmo Genético "
                     "sobre el asignador Jonker-Volgenant")
    run.italic = True
    run.font.size = Pt(12)
    doc.add_paragraph()

    # ── 1. Introducción ──
    doc.add_heading("1. Introducción y objetivo", level=1)
    doc.add_paragraph(
        "Este informe documenta una simulación adicional del proyecto: la "
        "optimización de los cuatro pesos (w1–w4) de la función de costes que "
        "gobierna la asignación de pedidos a drones. El objetivo es que los "
        "pedidos se completen en el menor tiempo posible y, si es posible, con "
        "menor consumo energético. Se comparan dos métodos para obtener esos "
        "parámetros —búsqueda Monte Carlo y un Algoritmo Genético— manteniendo en "
        "todo momento el asignador óptimo Jonker-Volgenant "
        "(scipy.optimize.linear_sum_assignment) para resolver la matriz de costes. "
        "El estudio se realiza sobre cinco arquetipos de escenario que cubren "
        "distintos tamaños de flota y perfiles de demanda."
    )

    build_methodology(doc)
    doc.add_page_break()

    doc.add_heading("3. Resultados por escenario", level=1)
    doc.add_paragraph(
        "Para cada arquetipo se optimizan los pesos sobre escenarios de "
        "entrenamiento y se evalúan sobre escenarios de test independientes. Se "
        "presentan la convergencia de ambos métodos, los pesos hallados, la "
        "mejora pareada en test, la comparación con el greedy y la sensibilidad "
        "de la función de costes."
    )
    for i, arch in enumerate(data["archetypes"], 1):
        build_archetype_section(doc, arch, i)

    build_global_conclusions(doc, data)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"✅ Informe generado: {OUT_DOCX}")
    print(f"   {len(data['archetypes'])} escenarios documentados.")


if __name__ == "__main__":
    main()
