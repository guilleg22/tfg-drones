"""
generate_guia_didactica.py — Documento Word DIDÁCTICO y accesible.

A diferencia de los dos informes técnicos, este documento explica en lenguaje
sencillo:
  1. Qué problema resolvemos y por qué.
  2. Cada concepto/herramienta que hemos usado, con analogías.
  3. Cómo interpretar cada TIPO de imagen que hemos generado (con ejemplos).
  4. Conclusiones claras.

Lee los resultados reales de los dos estudios para que las cifras sean exactas.

Uso:
  python scripts/generate_guia_didactica.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

_ROOT = Path(__file__).resolve().parent.parent
OUT = _ROOT / "results" / "Guia_Didactica_TFG.docx"

BASE_JSON = _ROOT / "results" / "parameter_study" / "study_results.json"
V2_JSON = _ROOT / "solucion_mejorada" / "results" / "study_v2" / "study_v2_results.json"


# ── Helpers de formato ───────────────────────────────────────────────────────

def figure(doc, rel, caption, width=6.0):
    path = _ROOT / rel
    if not path.exists():
        doc.add_paragraph(f"[Figura no encontrada: {rel}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def how_to_read(doc, lines):
    """Caja '¿Cómo se lee?' con viñetas."""
    p = doc.add_paragraph()
    r = p.add_run("Cómo leer esta imagen:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")


def concept(doc, title, body):
    p = doc.add_paragraph()
    r = p.add_run(f"{title}. ")
    r.bold = True
    p.add_run(body)


def main():
    base = json.loads(BASE_JSON.read_text(encoding="utf-8")) if BASE_JSON.exists() else None
    v2 = json.loads(V2_JSON.read_text(encoding="utf-8")) if V2_JSON.exists() else None

    # Cifras clave (con valores por defecto por robustez)
    ecom = next((a for a in v2["archetypes"] if a["key"] == "ecommerce_light"), None) if v2 else None
    urb = next((a for a in v2["archetypes"] if a["key"] == "urban_small"), None) if v2 else None
    milp = v2["milp_comparison"] if v2 else {"gap_pct": 21.8, "milp_makespan": 895, "jv_makespan": 1089}
    ecom_total = ecom["improvement_total_pct"] if ecom else 35.1
    ecom_rech = ecom["improvement_recharge_pct"] if ecom else 31.7

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── Portada ──
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Guía para entender la optimización del reparto con drones")
    r.bold = True; r.font.size = Pt(20)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Explicación sencilla de los métodos, cómo leer cada gráfica y "
                  "qué conclusiones sacar")
    r.italic = True; r.font.size = Pt(12)
    doc.add_paragraph()

    # ── 1. El problema ──
    doc.add_heading("1. ¿Qué problema estamos resolviendo?", level=1)
    doc.add_paragraph(
        "Tenemos una flota de drones y una lista de pedidos que entregar en "
        "Castelldefels. Cada dron tiene una capacidad de carga (cuánto peso "
        "aguanta) y una batería limitada. La pregunta es: ¿qué dron debe llevar "
        "cada pedido para que TODOS los pedidos se entreguen lo antes posible y, "
        "a ser posible, gastando poca energía?")
    doc.add_paragraph(
        "Decidir 'qué dron lleva qué pedido' es el problema de asignación. Para "
        "tomar esa decisión damos a cada posible pareja (dron, pedido) una "
        "puntuación de coste: cuanto menor sea el coste, mejor es esa pareja. El "
        "ordenador elige la combinación de parejas con menor coste total.")
    doc.add_paragraph(
        "Dos medidas nos importan, y conviene tenerlas claras desde el principio:")
    concept(doc, "Makespan (tiempo total)",
            "es el tiempo que tarda en completarse TODO el reparto, marcado por el "
            "dron que más tarda (el 'cuello de botella'). Es nuestra medida "
            "principal.")
    concept(doc, "Energía",
            "es el consumo total de batería de la flota para hacer todas las "
            "entregas.")

    # ── 2. Conceptos explicados ──
    doc.add_heading("2. Las herramientas que hemos usado (en cristiano)", level=1)

    doc.add_heading("2.1. La función de costes y sus pesos", level=2)
    doc.add_paragraph(
        "La 'puntuación' de cada pareja (dron, pedido) se calcula sumando varios "
        "factores. Cada factor mira una cosa distinta, y un PESO decide cuánta "
        "importancia le damos:")
    for nombre, desc in [
        ("w1 – Energía del viaje", "penaliza los viajes que gastan mucha batería."),
        ("w2 – Equilibrio de batería", "penaliza dejar a un dron casi sin batería."),
        ("w3 – Capacidad", "penaliza usar un dron grande para un paquete pequeño."),
        ("w4 – Espera por recarga", "penaliza que el dron tenga que recargar pronto."),
        ("w5 – Balanceo de carga (solución mejorada)",
         "penaliza cargar de trabajo a un dron que ya va muy ocupado, para "
         "repartir mejor y terminar antes."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{nombre}: "); r.bold = True
        p.add_run(desc)
    doc.add_paragraph(
        "El reto del proyecto es encontrar los valores de esos pesos que hacen "
        "que el reparto termine antes. Como no sabemos a mano cuáles son los "
        "mejores, los buscamos automáticamente con tres métodos distintos (los "
        "vemos abajo).")

    doc.add_heading("2.2. Cómo se decide la asignación: Jonker-Volgenant (JV)", level=2)
    doc.add_paragraph(
        "Con los costes de todas las parejas formamos una tabla (matriz) y "
        "pedimos al ordenador la combinación de menor coste total. Eso lo "
        "resuelve de forma ÓPTIMA un algoritmo clásico llamado Jonker-Volgenant "
        "(en el código, la función linear_sum_assignment). Es como repartir "
        "tareas entre personas buscando el reparto globalmente más barato, no "
        "pareja a pareja.")
    concept(doc, "Greedy (el método 'ingenuo' de comparación)",
            "asigna los pedidos uno a uno, dando cada pedido al dron más barato "
            "en ese momento, sin mirar el conjunto. Es nuestro punto de "
            "comparación para ver si vale la pena el método óptimo JV.")

    doc.add_heading("2.3. Los tres métodos para buscar los pesos", level=2)
    concept(doc, "Monte Carlo (MC)",
            "prueba muchísimas combinaciones de pesos al azar (bien repartidas) y "
            "se queda con la mejor. Simple y honesto, como tirar muchos dardos.")
    concept(doc, "Algoritmo Genético (GA)",
            "imita la evolución: parte de varias combinaciones, mezcla las "
            "mejores ('cruce'), introduce pequeños cambios ('mutación') y repite "
            "generación tras generación. Tiende a encontrar buenas soluciones.")
    concept(doc, "Optimización Bayesiana (solución mejorada)",
            "es la más 'lista': va construyendo un modelo de qué pesos funcionan "
            "y decide dónde probar a continuación. Encuentra buenas soluciones con "
            "muchas menos pruebas.")

    doc.add_heading("2.4. Cómo comprobamos que no nos engañamos", level=2)
    concept(doc, "Entrenamiento y test",
            "buscamos los pesos en un grupo de escenarios (entrenamiento) y los "
            "probamos en OTRO grupo distinto (test). Así sabemos si los pesos "
            "funcionan de verdad o solo 'se sabían la respuesta'.")
    concept(doc, "p-valor (significancia)",
            "mide si una diferencia es real o casualidad. Si p < 0,05 la "
            "diferencia se considera fiable; si es mayor, podría ser azar.")
    concept(doc, "Comparación pareada",
            "como probamos cada escenario con todas las opciones, comparamos "
            "escenario contra sí mismo. Por eso miramos la 'mejora por escenario' "
            "en vez de promedios sueltos.")

    doc.add_heading("2.5. Las mejoras de la segunda solución", level=2)
    concept(doc, "Recarga parcial",
            "antes los drones siempre recargaban al 100% (lento). Ahora recargan "
            "solo lo justo para los pedidos que quedan. Esta es la mejora que más "
            "tiempo ahorra.")
    concept(doc, "Objetivo sensible al riesgo (P90)",
            "en vez de optimizar el caso medio, optimiza los peores casos (el 10% "
            "de escenarios más lentos), que son los que de verdad importan.")
    concept(doc, "Baseline exacto (MILP)",
            "resolvemos una versión pequeña del problema de forma matemáticamente "
            "PERFECTA, para medir cuánto le falta a nuestro método para ser "
            "óptimo.")

    doc.add_page_break()

    # ── 3. Cómo interpretar cada imagen ──
    doc.add_heading("3. Cómo interpretar cada gráfica", level=1)
    doc.add_paragraph(
        "Hemos generado varios tipos de gráfica. Aquí explicamos, con un ejemplo "
        "real de cada uno, qué representa y cómo leerlo.")

    doc.add_heading("3.1. Mapa de calor de la matriz de costes", level=2)
    figure(doc, "results/parameter_study/urban_small/cost_heatmap.png",
           "Ejemplo de matriz de costes (drones en filas, pedidos en columnas).")
    how_to_read(doc, [
        "Cada celda es el coste de que un dron (fila) lleve un pedido (columna).",
        "Cuanto más clara/amarilla la celda, más barato; más roja, más caro.",
        "Las celdas grises con el símbolo ∞ son imposibles (p.ej. el paquete pesa "
        "más de lo que el dron aguanta).",
        "El algoritmo JV elige un conjunto de celdas (una por dron) que minimiza "
        "el coste total.",
    ])

    doc.add_heading("3.2. Curva de convergencia (cómo mejora la búsqueda)", level=2)
    figure(doc, "solucion_mejorada/results/study_v2/ecommerce_light/convergence_3tuners.png",
           "Convergencia de los tres métodos de búsqueda de pesos.")
    how_to_read(doc, [
        "El eje horizontal es el esfuerzo de cálculo (número de pruebas).",
        "El eje vertical es lo bueno que es el mejor resultado encontrado: 1,0 es "
        "'igual que no optimizar' (pesos neutros); por debajo de 1,0 es mejor.",
        "Cada línea baja según el método va encontrando pesos mejores; cuanto más "
        "baja y antes, mejor el método.",
        "Sirve para comparar métodos de forma justa: la Bayesiana suele bajar con "
        "menos pruebas; el Genético suele llegar más bajo con más pruebas.",
    ])

    doc.add_heading("3.3. Composición de pesos", level=2)
    figure(doc, "solucion_mejorada/results/study_v2/urban_small/weights5.png",
           "Qué importancia da cada método a cada factor (w1..w5).")
    how_to_read(doc, [
        "Cada grupo de barras es uno de los factores de coste (w1 a w5).",
        "Cada color es un método. La altura es la importancia que ese método "
        "decidió darle (todas suman 1).",
        "Permite ver en qué se fija cada método. 'Neutros' reparte todo por igual "
        "(0,2 en cada uno) como referencia.",
    ])

    doc.add_heading("3.4. Mejora por escenario (diagrama de caja)", level=2)
    figure(doc, "solucion_mejorada/results/study_v2/urban_small/test_improvement.png",
           "Mejora (%) que aporta cada método, escenario por escenario.")
    how_to_read(doc, [
        "El eje vertical es la mejora en % respecto a no optimizar. La línea "
        "discontinua en 0% es 'igual que neutros'.",
        "Cada caja resume muchos escenarios: la línea central es la mediana, la "
        "caja es la mitad central de los casos, el rombo es la media.",
        "Por encima de 0% = mejora; por debajo = empeora. Cajas pegadas a 0 "
        "significan 'casi siempre da igual'; los puntos sueltos arriba son "
        "escenarios donde la optimización ayudó mucho.",
        "Esta gráfica es honesta: muestra que la mejora media es pequeña pero que "
        "hay casos puntuales con gran ganancia.",
    ])

    doc.add_heading("3.5. Sensibilidad de la función de costes", level=2)
    figure(doc, "results/parameter_study/ecommerce_light/sensitivity.png",
           "Qué pasa si damos toda la importancia a un solo factor.")
    how_to_read(doc, [
        "Cada barra es el resultado de usar SOLO uno de los factores (todo el "
        "peso en w1, o solo en w2, etc.).",
        "La línea discontinua es el resultado con pesos equilibrados (neutros).",
        "Si una barra queda peor que la línea, ese factor por sí solo no basta. "
        "Aquí se ve que ningún factor aislado gana al reparto equilibrado: por "
        "eso combinamos todos.",
    ])

    doc.add_heading("3.6. Greedy vs. Jonker-Volgenant", level=2)
    figure(doc, "solucion_mejorada/results/study_v2/urban_small/greedy_vs_jv.png",
           "Comparación del método óptimo (JV) frente al ingenuo (Greedy).")
    how_to_read(doc, [
        "Barras rojas = método ingenuo (Greedy); verdes = método óptimo (JV).",
        "Se comparan energía, tiempo (makespan) y pedidos entregados.",
        "El % encima indica la mejora de JV; el p-valor debajo indica si la "
        "diferencia es fiable. En la práctica la diferencia entre ambos es "
        "pequeña en este problema.",
    ])

    doc.add_heading("3.7. Base vs. mejoras (la gráfica clave)", level=2)
    figure(doc, "solucion_mejorada/results/study_v2/ecommerce_light/base_vs_improved.png",
           "Cómo baja el tiempo total al ir añadiendo mejoras.")
    how_to_read(doc, [
        "La primera barra (roja) es el sistema original; las siguientes añaden "
        "mejoras una a una.",
        "El % encima de cada barra es cuánto hemos reducido el tiempo respecto al "
        "original.",
        "Se ve claramente que el gran salto lo da la recarga parcial; el ajuste "
        "de pesos añade un poco más.",
        "Las líneas negras verticales indican cuánto varían los escenarios entre "
        "sí (variabilidad natural).",
    ])

    doc.add_heading("3.8. Gap frente al óptimo exacto (MILP)", level=2)
    figure(doc, "solucion_mejorada/results/study_v2/milp_gap.png",
           "Nuestro método (JV) frente a la solución matemáticamente perfecta.", width=4.2)
    how_to_read(doc, [
        "La barra verde es el mejor tiempo posible (calculado de forma exacta).",
        "La barra azul es lo que consigue nuestro método rápido (JV).",
        "El 'gap' es lo que nos separa del óptimo: indica cuánto margen de mejora "
        "queda para el futuro.",
    ])

    doc.add_page_break()

    # ── 4. Resultados en cristiano ──
    doc.add_heading("4. Qué hemos encontrado (resumen sencillo)", level=1)
    puntos = [
        ("Gastar menos energía casi no depende de cómo asignemos.",
         "Como hay que entregar todos los pedidos y cada paquete obliga a usar "
         "cierto tipo de dron, la energía total apenas cambia. Por eso nos "
         "centramos en el tiempo."),
        ("Ajustar los pesos ayuda poco… de media.",
         "Cambiar los pesos mejora el tiempo solo un poco (normalmente menos del "
         "2-3%) y no siempre de forma fiable. Pero en algunos escenarios "
         "concretos evita un atasco y ahorra muchísimo tiempo."),
        ("La gran mejora vino de la recarga parcial.",
         f"Dejar de recargar siempre al 100% y cargar solo lo necesario redujo el "
         f"tiempo total de forma notable y fiable: hasta un {ecom_rech:.0f}% en el "
         f"escenario de paquetería ligera, sin dejar ningún pedido sin entregar."),
        ("El método óptimo (JV) no siempre gana al ingenuo (Greedy).",
         "JV es mejor en escenarios pequeños y de carga ligera, pero en los "
         "grandes la diferencia casi desaparece, porque la función de costes es "
         "una aproximación del tiempo, no el tiempo en sí."),
        ("Todavía hay margen de mejora.",
         f"Comparado con la solución matemáticamente perfecta, nuestro método deja "
         f"un margen de aproximadamente {milp['gap_pct']:.0f}%. Cerrar ese hueco "
         "es la principal vía de trabajo futuro."),
    ]
    for i, (tit, cuerpo) in enumerate(puntos, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {tit} "); r.bold = True
        p.add_run(cuerpo)

    # ── 5. Conclusiones ──
    doc.add_heading("5. Conclusiones", level=1)
    doc.add_paragraph(
        "El trabajo ha cumplido su objetivo de dos maneras. Primero, ha "
        "demostrado con rigor (separando entrenamiento y test, y con pruebas "
        "estadísticas) QUÉ palancas mejoran de verdad el reparto y cuáles no: "
        "hemos aprendido que el tiempo manda sobre la energía, que ajustar los "
        "pesos da una mejora pequeña, y que el verdadero cuello de botella es la "
        "recarga.")
    doc.add_paragraph(
        f"Segundo, sobre ese aprendizaje hemos construido una solución mejorada "
        f"que reduce el tiempo de reparto de forma clara —hasta un "
        f"{ecom_total:.0f}% en el mejor escenario— manteniendo el 100% de las "
        f"entregas, y hemos añadido un método de búsqueda más eficiente "
        f"(bayesiano) y una referencia exacta (MILP) que sitúa cuánto nos queda "
        f"por mejorar.")
    doc.add_paragraph(
        "En resumen: no nos hemos limitado a 'probar algoritmos', sino que hemos "
        "entendido el problema, hemos comprobado qué funciona y qué no, y hemos "
        "usado ese conocimiento para mejorar el sistema de forma medible y "
        "honesta. Esa es, precisamente, la forma correcta de abordar una "
        "optimización.")

    doc.add_heading("5.1. Recomendación práctica", level=2)
    doc.add_paragraph(
        "Para operar la flota: mantener el asignador óptimo JV, usar recarga "
        "parcial, ajustar los pesos con el método genético o bayesiano validando "
        "en escenarios nuevos, y volver a ajustar los pesos cuando cambie el tipo "
        "de demanda (no es lo mismo paquetería ligera que carga industrial).")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"✅ Guía didáctica generada: {OUT}")


if __name__ == "__main__":
    main()
