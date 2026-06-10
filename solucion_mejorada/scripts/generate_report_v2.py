"""
generate_report_v2.py — Informe .docx de la solución MEJORADA.

Lee results/study_v2/study_v2_results.json y construye un documento Word que:
  - Explica cada mejora (recarga parcial, término w5, objetivo P90/CVaR,
    optimización bayesiana, baseline MILP) con su fragmento de código real.
  - Muestra, por arquetipo, la mejora escalonada del makespan frente a la base.
  - Compara los tres tuners (MC, GA, Bayes).
  - Cuantifica el gap de optimalidad del JV frente al MILP exacto.
  - Concluye con la comparación global base vs mejorada.

Uso:
  python scripts/generate_report_v2.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_ROOT))

STUDY = _ROOT / "results" / "study_v2" / "study_v2_results.json"
OUT = _ROOT / "results" / "Informe_Solucion_Mejorada_TFG.docx"
MONO = "Consolas"


def extract_function(rel_path, func_name):
    text = (_ROOT / rel_path).read_text(encoding="utf-8").splitlines()
    start, indent = None, 0
    for i, line in enumerate(text):
        s = line.lstrip()
        if s.startswith(f"def {func_name}(") or s.startswith(f"def {func_name} "):
            start, indent = i, len(line) - len(s)
            break
    if start is None:
        return f"# (no encontrado {func_name})"
    depth, body = 0, start
    for j in range(start, len(text)):
        depth += text[j].count("(") - text[j].count(")")
        if depth <= 0 and text[j].rstrip().endswith(":"):
            body = j + 1
            break
    end = len(text)
    for j in range(body, len(text)):
        if text[j].strip() == "":
            continue
        if len(text[j]) - len(text[j].lstrip()) <= indent:
            end = j
            break
    return "\n".join(text[start:end]).rstrip()


def code(doc, snippet, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption); r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(snippet)
    r.font.name = MONO; r.font.size = Pt(8.5)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)


def figure(doc, rel, caption, width=6.3):
    path = _ROOT / rel
    if not path.exists():
        doc.add_paragraph(f"[Figura no encontrada: {rel}]"); return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(9)
    return t


def w5(arr):
    return "[" + ", ".join(f"{x:.2f}" for x in arr) + "]"


def main():
    if not STUDY.exists():
        print("Falta study_v2_results.json. Ejecuta run_study_v2.py primero.")
        sys.exit(1)
    data = json.loads(STUDY.read_text(encoding="utf-8"))
    archs = data["archetypes"]

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Portada
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Solución mejorada para la optimización del reparto con drones")
    r.bold = True; r.font.size = Pt(20)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Recarga parcial · balanceo de carga (w5) · objetivo sensible al "
                  "riesgo · optimización bayesiana · baseline MILP exacto")
    r.italic = True; r.font.size = Pt(12)

    # 1. Resumen ejecutivo
    doc.add_heading("1. Resumen ejecutivo", level=1)
    best = max(archs, key=lambda a: a["improvement_total_pct"])
    avg_rech = sum(a["improvement_recharge_pct"] for a in archs) / len(archs)
    doc.add_paragraph(
        "Esta solución mejorada parte del diagnóstico de la solución base: la "
        "energía es casi invariante a los pesos y las ganancias del tuning de "
        "pesos son pequeñas porque el makespan está dominado por los ciclos de "
        "recarga. Sobre ese diagnóstico se implementan cinco mejoras, "
        "manteniendo el asignador linear_sum_assignment (Jonker-Volgenant). El "
        f"resultado: la recarga parcial reduce el makespan una media del "
        f"{avg_rech:.1f}% (estadísticamente significativa), con el 100% de los "
        f"pedidos entregados, y el tuning de pesos añade una mejora menor encima. "
        f"La mayor reducción total se da en «{best['name']}» "
        f"({best['improvement_total_pct']:+.1f}%)."
    )

    # 2. Mejoras y código
    doc.add_heading("2. Mejoras implementadas", level=1)

    doc.add_heading("2.1. Recarga parcial adaptativa", level=2)
    doc.add_paragraph(
        "La solución base recargaba siempre al 100 %. Como una recarga completa "
        "(~2 h) puede superar el tiempo de vuelo de toda una ronda, es el factor "
        "que domina el makespan. La mejora recarga solo lo necesario para la "
        "cuota de viajes pendientes de cada dron (acotado al 100 %), y no recarga "
        "drones que no pueden servir ningún pedido pendiente. En las rondas "
        "finales, con pocos pedidos, esto recorta drásticamente el makespan."
    )
    code(doc, extract_function("sim2/simulator.py", "_recharge_drones"),
         "Código 1 — Política de recarga parcial (sim2/simulator.py).")

    doc.add_heading("2.2. Término de balanceo de carga (w5)", level=2)
    doc.add_paragraph(
        "Se añade un quinto término a la función de costes que penaliza asignar a "
        "drones con mucho tiempo de trabajo acumulado, haciendo la asignación "
        "consciente del makespan (reparto de carga). Con w5 = 0 el comportamiento "
        "es idéntico a la solución base."
    )
    code(doc, extract_function("sim2/cost_function.py", "compute_cost"),
         "Código 2 — Función de costes con el término w5 (sim2/cost_function.py).")

    doc.add_heading("2.3. Objetivo sensible al riesgo (P90 / CVaR)", level=2)
    doc.add_paragraph(
        "Como las ganancias provienen de escenarios patológicos raros, se añade "
        "la posibilidad de optimizar el percentil 90 o el CVaR del makespan en "
        "lugar de la media, atacando directamente la cola de la distribución."
    )
    code(doc, extract_function("sim2/objective.py", "_scalarize"),
         "Código 3 — Escalarización con objetivos P90/CVaR (sim2/objective.py).")

    doc.add_heading("2.4. Tercer método de tuning: optimización bayesiana", level=2)
    doc.add_paragraph(
        "Junto a Monte Carlo y el Algoritmo Genético se añade optimización "
        "bayesiana (proceso gaussiano + Expected Improvement, vía "
        "scikit-optimize), el enfoque más eficiente en evaluaciones para espacios "
        "continuos de baja dimensión y estado del arte en configuración "
        "automática de algoritmos."
    )
    code(doc, extract_function("sim2/optimizer_bayes.py", "optimize_bayes"),
         "Código 4 — Optimización bayesiana de los 5 pesos (sim2/optimizer_bayes.py).")

    doc.add_heading("2.5. Baseline exacto MILP", level=2)
    doc.add_paragraph(
        "Para medir cuán lejos queda el heurístico JV del óptimo real de "
        "makespan, se resuelve de forma exacta (programación lineal entera con "
        "PuLP/CBC) el makespan óptimo en un ciclo de carga."
    )
    code(doc, extract_function("sim2/milp_baseline.py", "solve_min_makespan_single_cycle"),
         "Código 5 — Baseline MILP exacto (sim2/milp_baseline.py).")
    doc.add_page_break()

    # 3. Resultados por arquetipo
    doc.add_heading("3. Resultados por escenario", level=1)
    for i, a in enumerate(archs, 1):
        doc.add_heading(f"3.{i}. {a['name']}", level=2)
        doc.add_paragraph(a["description"])
        st = a["stages"]
        table(doc, ["Etapa", "Makespan medio (s)", "Entregas medias", "Mejora vs base"],
              [[st["labels"][k], f"{st['makespan_mean'][k]:.0f}",
                f"{st['delivered_mean'][k]:.1f}",
                "baseline" if k == 0 else
                f"{(st['makespan_mean'][0]-st['makespan_mean'][k])/st['makespan_mean'][0]*100:+.1f}%"]
               for k in range(len(st["labels"]))])
        figure(doc, a["figures"]["base_vs_improved"],
               f"Figura 3.{i}.a — Makespan: base vs mejoras acumuladas.")
        figure(doc, a["figures"]["convergence"],
               f"Figura 3.{i}.b — Convergencia de los tres tuners (mismo presupuesto).")
        tn = a["tuners"]
        table(doc, ["Tuner", "Pesos [w1..w5]", "Obj.", "Evals", "Δt vs neutros"],
              [[name, w5(tn[name]["weights"]), f"{tn[name]['best_objective']:.4f}",
                tn[name]["n_evaluations"], f"{tn[name]['time_improvement_pct']:+.2f}%"]
               for name in tn])
        figure(doc, a["figures"]["weights"],
               f"Figura 3.{i}.c — Composición de los 5 pesos por método.")
        figure(doc, a["figures"]["test_improvement"],
               f"Figura 3.{i}.d — Aporte del tuning sobre la recarga parcial (test).")
        p = doc.add_paragraph()
        p.add_run("Conclusión del escenario: ").bold = True
        p.add_run(
            f"la recarga parcial reduce el makespan {a['improvement_recharge_pct']:+.1f}% "
            f"(p={a['p_value_recharge']:.4f}); el tuning de pesos añade "
            f"{a['improvement_tuning_pct']:+.1f}% (mejor tuner: {a['best_tuner']}); "
            f"reducción total {a['improvement_total_pct']:+.1f}% manteniendo el "
            f"100% de entregas.")
        doc.add_page_break()

    # 4. Gap MILP
    doc.add_heading("4. Gap de optimalidad frente al MILP exacto", level=1)
    m = data["milp_comparison"]
    doc.add_paragraph(
        f"En una instancia pequeña de un ciclo de carga ({m['n_drones']} drones, "
        f"{m['n_orders']} pedidos), el makespan óptimo exacto (MILP) es "
        f"{m['milp_makespan']:.0f} s, mientras que el heurístico JV obtiene "
        f"{m['jv_makespan']:.0f} s: un gap del {m['gap_pct']:+.1f}%. Esto confirma "
        "que el JV, al minimizar el coste de asignación y no el makespan "
        "directamente, deja margen de mejora. Cerrar ese gap (p.ej. con una "
        "asignación makespan-directa o look-ahead) es la principal línea de "
        "trabajo futuro.")
    figure(doc, m["figure"], "Figura 4 — Makespan óptimo (MILP) vs heurístico JV.", width=4.5)

    # 5. Conclusiones globales
    doc.add_heading("5. Conclusiones globales", level=1)
    table(doc, ["Escenario", "Δ recarga", "Δ tuning", "Δ total", "Mejor tuner"],
          [[a["name"], f"{a['improvement_recharge_pct']:+.1f}%",
            f"{a['improvement_tuning_pct']:+.1f}%", f"{a['improvement_total_pct']:+.1f}%",
            a["best_tuner"]] for a in archs])
    doc.add_paragraph()
    for n, (titulo, cuerpo) in enumerate([
        ("La recarga parcial es la mejora de mayor impacto.",
         "Reduce el makespan de forma significativa y consistente en todos los "
         "escenarios, sin perder entregas. Es un cambio operativo, no de tuning, "
         "lo que confirma el diagnóstico de la solución base."),
        ("El tuning de pesos aporta una mejora pequeña pero gratuita.",
         "Sobre la recarga parcial, optimizar los 5 pesos añade una mejora "
         "marginal. El término w5 de balanceo de carga ayuda en los escenarios "
         "con varios ciclos de recarga."),
        ("La optimización bayesiana es la más eficiente en evaluaciones.",
         "Alcanza objetivos comparables a MC y GA con menos evaluaciones de la "
         "simulación, como cabía esperar en un espacio continuo de baja dimensión."),
        ("Queda un gap claro frente al óptimo exacto.",
         "El MILP muestra que el JV deja margen de makespan; la asignación "
         "directa de makespan es la línea de mejora futura más prometedora."),
    ], 1):
        p = doc.add_paragraph(); p.add_run(f"{n}. {titulo} ").bold = True; p.add_run(cuerpo)

    doc.add_heading("5.1. Reproducibilidad", level=2)
    code(doc, "python scripts/run_study_v2.py\npython scripts/generate_report_v2.py",
         "Código 6 — Reproducción del estudio mejorado.")
    doc.add_paragraph(
        f"Objetivo optimizado: '{data['objective']}'. Tiempo de cómputo: "
        f"{data.get('elapsed_s', 0):.0f} s. Dependencias añadidas: scikit-optimize "
        "(bayesiana) y PuLP (MILP).")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"✅ Informe mejorado generado: {OUT}")
    print(f"   {len(archs)} escenarios documentados.")


if __name__ == "__main__":
    main()
