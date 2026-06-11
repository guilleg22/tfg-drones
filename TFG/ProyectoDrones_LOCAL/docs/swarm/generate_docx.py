"""
Genera el documento Word "TFG_Swarm_Manager.docx" con toda la planificacion
de la entrega del enjambre de drones, bibliografia incluida.

Uso:
    python generate_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent / "TFG_Swarm_Manager.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(it)
        run.font.size = Pt(11)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Sombrear el fondo via shd al parrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table_simple(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], "1F3A5F")

    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_reference(doc, n, authors, year, title, source, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    run = p.add_run("[{}] ".format(n))
    run.bold = True
    run.font.size = Pt(10)
    text = "{}. ({}). {} ".format(authors, year, title)
    if source:
        text += "{}. ".format(source)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    add_hyperlink(p, url, url)


# ---------------------------------------------------------------------------
# Documento
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------------- PORTADA ----------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Trabajo Fin de Grado")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("Gestion autonoma de un enjambre de drones\n"
                         "para reparto de paquetes en Mission Planner")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Documento de planificacion de la entrega Swarm Manager")
    run.italic = True
    run.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Proyecto: ProyectoDrones_LOCAL  -  Desktop Drone Control v2.0\n").italic = True
    info.add_run("Algoritmo de planificacion y orquestacion para flota heterogenea de drones autonomos\n").italic = True
    info.add_run("Fecha del documento: 2026-05-21\n").italic = True
    info.add_run("Autor: Guillermo Galve\n").italic = True

    doc.add_page_break()

    # ---------------- INDICE / RESUMEN ----------------
    add_heading(doc, "Resumen ejecutivo", level=1)
    add_para(doc,
        "El proyecto Desktop Drone Control v2.0 permite operar un unico dron contra "
        "Mission Planner / SITL mediante la libreria dronLink y asigna a cada pedido "
        "la ruta predefinida mas cercana al cliente. La siguiente entrega del TFG debe "
        "escalar el sistema a un enjambre heterogeneo de drones que parten todos de "
        "un parking, disponen de rutas pre-aprobadas, distintas capacidades de carga, "
        "diferentes niveles de bateria y autonomia, y un peso por pedido. Un Swarm "
        "Manager debera decidir de forma autonoma y optima que dron coge que pedidos, "
        "en que orden y por que corredor, evitando conflictos en el espacio aereo y "
        "maximizando la utilizacion de la flota.")
    add_para(doc,
        "Este documento condensa toda la planificacion: auditoria del estado actual, "
        "estado del arte revisado y citado, arquitectura propuesta, roadmap por fases y "
        "lista granular de tareas. Es el documento maestro que precede a la "
        "implementacion y que servira como anexo del documento del TFG.")

    add_heading(doc, "Alcance del trabajo", level=2)
    add_para(doc, "Dentro del alcance:", bold=True)
    add_bullets(doc, [
        "Modelado de la flota: dron con id, autonomia (Wh o min), payload maximo, estado y posicion.",
        "Algoritmo de asignacion optimo (o casi-optimo) considerando peso del pedido, autonomia restante, balanceo de carga y tiempo total de entrega (makespan).",
        "Orquestador de ejecucion: traduccion del plan a comandos MAVLink concurrentes (un hilo por dron, instancias SITL independientes).",
        "Deconfliccion basica: separacion por slot de altitud y secuenciacion de despegues para evitar colisiones en el HUB.",
        "Visualizacion en el dashboard PySide6: panel de flota multi-dron, mapa con todos los marcadores y trayectorias activas.",
        "Validacion en SITL con minimo 3 drones simulados y minimo 5 pedidos entrantes.",
    ])
    add_para(doc, "Fuera del alcance:", bold=True)
    add_bullets(doc, [
        "Detect-and-avoid reactivo sensor-based: solo se aborda deconfliccion estrategica.",
        "Estaciones de recarga fisicamente distribuidas: solo se simula tiempo de recarga.",
        "Aprendizaje por refuerzo y redes neuronales: se utiliza optimizacion clasica (CVRP + metaheuristica) que es defendible y reproducible.",
        "Hardware real: la entrega se valida integramente sobre SITL.",
    ])

    doc.add_page_break()

    # ---------------- ESTADO ACTUAL ----------------
    add_heading(doc, "1. Estado actual del proyecto", level=1)
    add_para(doc,
        "El proyecto esta organizado en capas limpias (GUI con PySide6, servicios, "
        "capa de negocio con SQLite y un servidor HTTP para el portal cliente) pero "
        "asume un unico dron identificado por la constante DRONE_ID = 'Dron-1' en "
        "cliente/main_window.py:24. Toda la logica de entrega vive en "
        "DroneService.start_order_delivery (servicios/drone_service.py:298) y se "
        "ejecuta secuencialmente para un solo autopiloto conectado por TCP al puerto "
        "5763 del Mission Planner / SITL local.")

    add_heading(doc, "1.1 Modelo de datos actual", level=2)
    add_para(doc,
        "La base de datos operations.db contiene dos tablas: clients(id, name, "
        "address, latitude, longitude) y orders(id, client_id, weight_kg, status, "
        "assigned_profile_name, assigned_route_name, assigned_destination_*, "
        "assigned_distance_km, operational_state). No existe ninguna tabla drones; "
        "todo el sistema asume implicitamente un solo vehiculo. Los estados validos "
        "de orders.status son: pendiente, planificado, en_reparto, entregado, "
        "cancelado.")
    add_para(doc,
        "El fichero route_profiles.json define perfiles geograficos con hub central, "
        "parkings, destinos y rutas pre-aprobadas con waypoints intermedios. Cada "
        "ruta es ya un corredor aereo valido, lo que simplifica enormemente la "
        "deconfliccion porque el espacio aereo esta pre-aprobado; el Swarm Manager "
        "solo debe ordenar quien entra a cada corredor y a que altitud.")

    add_heading(doc, "1.2 Flujo de un pedido (mono-dron)", level=2)
    add_numbered(doc, [
        "El cliente entra en el portal web (servidor/api_server.py en :8080) y crea un pedido con peso.",
        "POST /api/orders invoca DeliveryDataStore.create_order que asigna el destino mas cercano por Haversine.",
        "El polling en MainWindow._poll_active_order() refresca el panel cada 2 segundos con los pedidos pendientes.",
        "El operador pulsa 'Aceptar' y MainWindow._on_accept_order() cambia el estado a en_reparto y llama a drone_svc.start_order_delivery().",
        "DroneService.start_order_delivery ejecuta secuencialmente: takeoff 25m -> goto parking -> goto hub -> intermedios -> destino -> cliente -> Land -> espera 10s -> re-armado -> takeoff -> goto central -> Land.",
    ])

    add_heading(doc, "1.3 Limitaciones para escalar a enjambre", level=2)
    rows = [
        ("DRONE_ID hardcoded en main_window.py:24", "Toda la UI asume un dron. Hay que parametrizar."),
        ("DroneService mantiene un solo Dron()", "No se puede orquestar mas de un autopiloto."),
        ("Mission Planner en 127.0.0.1:5763 unico", "Multi-SITL requiere un puerto por dron."),
        ("start_order_delivery() es bloqueante por dron", "OK por hilo, pero no decide quien la coge."),
        ("Matching pedido->ruta solo por distancia", "No considera peso, bateria ni dron concreto."),
        ("Sin tabla drones, sin capacidad ni bateria", "Hay que crear todo el modelo de flota."),
        ("Sin cola autonoma de pedidos", "El swarm debe planificar sin operador humano."),
        ("Sin deconfliccion entre rutas", "Dos drones con rutas que comparten hub -> conflicto."),
    ]
    add_table_simple(doc,
        headers=["Problema", "Consecuencia"],
        rows=rows, col_widths=[7, 10])

    add_heading(doc, "1.4 Puntos de extension naturales", level=2)
    add_para(doc,
        "El codigo esta bien separado en capas, lo que facilita la extension sin "
        "reescribir lo existente. Los enganches naturales son:")
    add_bullets(doc, [
        "servicios/ es el lugar para swarm_service.py (orquestador) y planner_service.py (motor VRP).",
        "negocio/db_manager.py absorbe las nuevas tablas drones y assignments.",
        "modelos/ se amplia con drone.py, assignment.py y flight_plan.py.",
        "widgets/fleet_panel.py ya tiene add_drone(drone_id) para N tarjetas.",
        "cliente/main_window.py mantiene la idea de un DroneService; se generaliza a dict[str, DroneService].",
        "servidor/api_server.py mantiene POST /api/orders; los pedidos se encolan en lugar de asignarse estaticamente.",
    ])

    doc.add_page_break()

    # ---------------- ESTADO DEL ARTE ----------------
    add_heading(doc, "2. Estado del arte y marco teorico", level=1)
    add_para(doc,
        "El problema de gestionar un enjambre de N drones que parten de un parking "
        "y deben entregar M paquetes con peso y rutas pre-aprobadas cae en la familia "
        "de los Vehicle Routing Problems with Drones (VRPD). Mas concretamente, "
        "combina varias variantes clasicas de la literatura de investigacion operativa "
        "y logistica.")

    add_heading(doc, "2.1 Formulacion del problema", level=2)
    add_table_simple(doc,
        headers=["Variante", "Aplica", "Justificacion"],
        rows=[
            ("CVRP (Capacitated VRP)", "Si", "Cada dron tiene capacidad de carga limitada."),
            ("HFVRP (Heterogeneous Fleet)", "Si", "Drones con distinta autonomia y categoria de carga."),
            ("E-VRP (Energy-constrained)", "Si", "Bateria finita; consumo depende de payload y distancia."),
            ("MDVRP (Multi-Depot)", "Parcial", "Un unico parking en el caso base; el codigo ya soporta varios."),
            ("VRPTW (Time Windows)", "Futuro", "Si el TFG anade SLA por pedido en una iteracion posterior."),
        ],
        col_widths=[4, 2, 11])
    add_para(doc,
        "La formulacion canonica del Drone Routing Problem la introducen Dorling et al. "
        "(2017) [1] y ha sido ampliada de forma consistente desde entonces, con revisiones "
        "recientes como la de Liu et al. (2023) [9].")

    add_heading(doc, "2.2 Algoritmos de asignacion y enrutado", level=2)
    add_para(doc, "Solvers centralizados exactos:", bold=True)
    add_para(doc,
        "Google OR-Tools [5] proporciona una API Python madura, gratuita y sin "
        "compilacion nativa en Windows, que soporta restricciones de capacidad "
        "(AddDimensionWithVehicleCapacity) y de distancia/tiempo (AddDimension). "
        "La estrategia recomendada es First-Solution + Guided Local Search con "
        "un timeout de pocos segundos. Es la opcion por defecto adoptada en este "
        "TFG por su madurez y por la documentacion disponible.")
    add_para(doc,
        "Como baseline simple, la asignacion pura uno-a-uno se puede resolver con el "
        "algoritmo Hungarian en O(n^3) mediante scipy.optimize.linear_sum_assignment. "
        "Es util como fallback si OR-Tools no encuentra solucion factible.")
    add_para(doc, "Algoritmos descentralizados (referencia teorica):", bold=True)
    add_para(doc,
        "El Consensus-Based Bundle Algorithm (CBBA) introducido por Choi et al. "
        "(2009) [3] permite asignacion descentralizada por puja en mercado entre "
        "agentes con consenso local. Trabajos recientes como Two-Level Clustered "
        "CBBA (2025) [11] o las extensiones para entornos dinamicos demuestran su "
        "viabilidad para enjambres heterogeneos. En este TFG el CBBA se cita como "
        "alternativa para trabajo futuro: nuestro Mission Planner es centralizado.")
    add_para(doc, "Metaheuristicas (mejora opcional):", bold=True)
    add_para(doc,
        "Existen variantes basadas en algoritmos geneticos con rescheduling [8] o "
        "estrategias hibridas multi-objetivo. Para los tamanos del TFG (<=10 drones, "
        "<=30 pedidos) la metaheuristica integrada en OR-Tools (Guided Local Search) "
        "es suficiente y no se anade una propia.")

    add_heading(doc, "2.3 Modelo de energia", level=2)
    add_para(doc,
        "La autonomia es la restriccion operativa mas relevante en drones de reparto. "
        "Para este TFG se adopta un modelo lineal validado en la literatura, donde la "
        "energia consumida es aproximadamente proporcional al tiempo de vuelo y al "
        "peso del payload [12]:")
    add_code(doc,
        "E_consumida = alpha * t_vuelo + beta * m_payload * d_horizontal + gamma * d_vertical")
    add_para(doc,
        "Los coeficientes alpha, beta y gamma se calibran experimentalmente y se "
        "consideran lineales en payload, segun Stolaroff et al. (2018) y reproducido "
        "por Zhang et al. (2021) [12]. Estudios alternativos sobre escenarios con viento "
        "variable son aportados por Bryant et al. (2020) [2]. Si la energia estimada para "
        "el trayecto completo (ida + retorno al parking) supera la bateria disponible con "
        "margen de seguridad superior al 20 por ciento, el dron no puede coger ese pedido.")
    add_para(doc,
        "El modelo battery-aware no lineal de Abeywickrama et al. (2018) [1bat] considera "
        "que la capacidad efectiva de la bateria se degrada al disminuir el SOC. Solo merece "
        "la pena implementarlo si se extiende al modelo de Peukert; queda fuera del alcance "
        "del TFG. Adaptaciones modernas basadas en RL como A-POMO (AIAA, 2024) [12rl] o los "
        "trabajos de Lakhdari et al. [6][7] se citan como referencia teorica.")

    add_heading(doc, "2.4 Coordinacion multi-vehiculo en ArduPilot / MAVLink", level=2)
    add_para(doc,
        "La configuracion fundamental para volar varios drones en ArduPilot es asignar "
        "un SYSID_THISMAV unico (1..255) a cada autopiloto [15]. En SITL esto se logra "
        "con sim_vehicle.py --instance N --auto-sysid, abriendo cada instancia un TCP "
        "diferente (5760, 5763, 5766, ...). Mission Planner soporta swarming en beta [14] "
        "abriendo varias conexiones MAVLink simultaneas y reenviando posiciones GPS del "
        "lider a seguidores en modo GUIDED.")
    add_para(doc,
        "El protocolo MAVLink Mission Protocol [16] define UPLOAD, DOWNLOAD, SET_CURRENT "
        "y MISSION_ACK. Al subir misiones a N drones en paralelo el ACK puede tardar; se "
        "debe hacer en threads independientes y no serializar. El framework DroneKit-Python "
        "y pymavlink son las opciones Python establecidas; el proyecto actual usa dronLink, "
        "que es una capa propia construida sobre pymavlink.")

    add_heading(doc, "2.5 Gestion del espacio aereo (UTM)", level=2)
    add_para(doc,
        "Para v1 solo se aborda deconfliccion estrategica (en el plan, antes de volar). "
        "El marco teorico UTM define tres capas de gestion: deconfliccion estrategica, "
        "deconfliccion tactica y detect-and-avoid [17]. Estudios como Aerospace MDPI (2023) [13a] "
        "proponen estructuras de espacio aereo en corredores, lanes y tubes para garantizar "
        "separaciones laterales, verticales y longitudinales seguras.")
    add_para(doc,
        "La estrategia aplicada en este TFG combina dos mecanismos sencillos pero efectivos:")
    add_bullets(doc, [
        "Slot de altitud por dron en vuelo: capas a 25, 30, 35, 40 y 45 metros. El planner asigna a cada dron una altitud de crucero distinta de las que estan ya en vuelo.",
        "Escalonado de despegues: si dos drones tienen que salir del mismo parking en t ~= 0, se introduce un retraso TAKEOFF_GAP_S = 10 segundos entre ambos.",
        "El HUB es un punto sensible: el planner garantiza que como maximo un dron lo ocupe en un instante dado calculando ETAs.",
        "Opcional: geofence virtual del parking con admision de un solo dron en la fase de descenso.",
    ])

    add_heading(doc, "2.6 Resumen de decisiones del TFG con justificacion", level=2)
    add_table_simple(doc,
        headers=["Decision", "Opcion elegida", "Justificacion"],
        rows=[
            ("Algoritmo principal", "CVRP heterogeneo + energia via OR-Tools",
             "Madurez, soporte Python, suficiente para tamanos del TFG."),
            ("Algoritmo fallback", "Hungarian (scipy)",
             "Trivial, O(n^3), siempre devuelve algo si OR-Tools falla."),
            ("Modelo de energia", "Lineal (Stolaroff/Zhang)",
             "Apto para optimizacion, justificacion bibliografica solida."),
            ("Arquitectura comms", "Centralizada (un GCS)",
             "Cuadra con Desktop Drone Control. CBBA citado como trabajo futuro."),
            ("MAVLink multi-dron", "TCP por instancia SITL",
             "Documentado en ArduPilot oficial; encaja con dronLink."),
            ("Deconfliccion", "Layered altitudes + escalonado",
             "Barata, demostrable visualmente; corredores pre-aprobados."),
            ("UI", "Extension de FleetPanel",
             "add_drone ya existe; basta con cablear N drones."),
            ("Persistencia", "Ampliacion SQLite",
             "Minimo coste, no introduce dependencias nuevas."),
        ], col_widths=[3.5, 5, 8])

    doc.add_page_break()

    # ---------------- ARQUITECTURA ----------------
    add_heading(doc, "3. Arquitectura propuesta", level=1)
    add_para(doc,
        "El modulo Swarm Manager se construye encima del proyecto actual sin reescribir "
        "lo existente. Se anaden ficheros nuevos en las capas modelos/, servicios/ y "
        "negocio/, ademas de una carpeta config/ con la declaracion de la flota.")

    add_heading(doc, "3.1 Vista de alto nivel", level=2)
    add_code(doc,
"""+----------------------------------------------------------------+
|                     UI (PySide6)                                |
|  MainWindow <-> FleetPanel <-> DroneCard (N) <-> MapWidget      |
+----------------------------------------------------------------+
                              ^
                              | senales Qt
                              v
+----------------------------------------------------------------+
|                  SwarmService (orquestador)                     |
|  [FleetState] [Planner CVRP] [AirspaceManager] [Executor (Nx)]  |
+----------------------------------------------------------------+
                              ^
                              | comandos MAVLink (paralelo)
                              v
+--------------------+   +--------------------+
|  DroneService(D1)  |   |  DroneService(D2)  |  ...
|  dronLink->:5760   |   |  dronLink->:5763   |
+--------------------+   +--------------------+
                              ^
                              | TCP (un puerto por dron)
                              v
                Mission Planner / SITL multi-instancia
                     (SYSID_THISMAV unicos)
""")

    add_heading(doc, "3.2 Nuevo arbol de paquetes", level=2)
    add_code(doc,
"""ProyectoDrones_LOCAL/
  modelos/
    waypoint.py            (existente)
    drone.py               * NUEVO  - DTO de dron, bateria, perfil
    assignment.py          * NUEVO  - Pedido <-> Dron + ruta + ETA
    flight_plan.py         * NUEVO  - Plan completo (varios assignments)
  negocio/
    db_manager.py          (modificado: + tablas drones, assignments)
    fleet_repository.py    * NUEVO  - CRUD de drones
  servicios/
    drone_service.py       (modificado: drone_id en senales)
    swarm_service.py       * NUEVO  - orquestador global
    planner_service.py     * NUEVO  - wrapper OR-Tools (CVRP)
    energy_model.py        * NUEVO  - consumo lineal
    airspace_manager.py    * NUEVO  - slots de altitud + escalonado
    executor_service.py    * NUEVO  - ejecuta un assignment en un dron
  widgets/
    fleet_panel.py         (modificado: N tarjetas + estado por dron)
  config/
    fleet.yaml             * NUEVO  - definicion declarativa de la flota
""")

    add_heading(doc, "3.3 Modelo de dominio", level=2)
    add_para(doc,
        "modelos/drone.py define tres dataclasses: DroneSpec (especificacion estatica "
        "con id, sysid, conexion TCP, payload maximo, capacidad de bateria, velocidades "
        "y coeficientes de consumo), DroneRuntime (estado mutable: bateria actual, "
        "lat/lon/alt, heading, slot de altitud asignado, assignment activo) y Drone "
        "(agrupa spec + runtime con metodo can_carry()).")
    add_para(doc,
        "modelos/assignment.py define Assignment con drone_id, order_id, perfil y ruta "
        "elegidos, slot de altitud, takeoff_offset_s, distancia y energia estimadas, "
        "y status (planned/executing/done/aborted).")
    add_para(doc,
        "modelos/flight_plan.py define FlightPlan como contenedor de un ciclo del "
        "planner: lista de assignments, pedidos no asignables en este ciclo y el coste "
        "del solver (segundos o Wh, segun metrica configurada).")

    add_heading(doc, "3.4 Persistencia ampliada", level=2)
    add_para(doc,
        "Se crean cuatro tablas SQLite nuevas en operations.db:")
    add_bullets(doc, [
        "drones: catalogo estatico (drone_id PK, sysid UNIQUE, connection, max_payload_kg, battery_capacity_wh, cruise_speed_mps, coeficientes de consumo, parking de origen, enabled).",
        "drone_states: estado en tiempo real (drone_id PK, state, battery_pct, battery_wh, lat/lon/alt, cruise_alt_slot, current_assignment_id, updated_at).",
        "flight_plans: historico de planes generados (plan_id PK, created_at, solver_objective, unassigned_orders en JSON).",
        "assignments: relacion plan-dron-pedido (assignment_id PK, plan_id, drone_id, order_id, profile/route, cruise_alt_m, takeoff_offset_s, distancia/duracion/energia estimadas, status, started_at, finished_at).",
    ])
    add_para(doc,
        "Se introduce el estado 'asignado' en orders.status, entre pendiente y "
        "en_reparto, para reflejar que el planner ya lo metio en un plan pero el "
        "dron aun no ha despegado. La definicion de la flota se hace de forma "
        "declarativa en config/fleet.yaml para reproducibilidad.")

    add_heading(doc, "3.5 Servicios nuevos", level=2)
    add_para(doc, "energy_model.py", bold=True)
    add_para(doc,
        "Implementa el modelo lineal de la seccion 2.3 con la funcion pura "
        "estimate_energy_wh(drone_spec, route_profile, weight_kg, distance_km, "
        "altitude_gain_m). Test unitario obligatorio con casos calibrados a mano.")
    add_para(doc, "planner_service.py", bold=True)
    add_para(doc,
        "Dado el estado de la flota, los pedidos pendientes y el catalogo de rutas, "
        "devuelve un FlightPlan. El algoritmo construye nodos (0 = parking, "
        "1..M = pedidos), calcula para cada par (drone, pedido) la ruta candidata, "
        "distancia, energia y viabilidad, define un RoutingModel de OR-Tools con "
        "dimensiones de capacidad (peso) y energia (Wh con margen 20%), aplica "
        "PATH_CHEAPEST_ARC + GUIDED_LOCAL_SEARCH con timeout de 5s y mapea la "
        "solucion a Assignments. Si la solucion es vacia, hace fallback a "
        "Hungarian via scipy.optimize.linear_sum_assignment.")
    add_para(doc, "airspace_manager.py", bold=True)
    add_para(doc,
        "Mantiene el conjunto de slots de altitud libres (25, 30, 35, 40 y 45 m por "
        "defecto), asigna slot a cada Assignment saliente, libera el slot cuando el "
        "dron regresa al parking y calcula takeoff_offset_s respetando una separacion "
        "minima TAKEOFF_GAP_S = 10s. Garantiza ademas que no haya mas de un dron en "
        "el HUB en un instante dado.")
    add_para(doc, "executor_service.py", bold=True)
    add_para(doc,
        "Encapsula la ejecucion de un Assignment sobre un DroneService. Reproduce el "
        "flujo actual de start_order_delivery pero parametrizado: esperar takeoff_offset, "
        "armar y despegar al cruise_alt_m del slot, goto parking-hub-intermedios-destino-"
        "cliente, Land, sleep entrega, re-arm, takeoff retorno, goto cliente-destino-hub-"
        "parking, Land. Emite senales Qt progress / finished / failed. Cada dron tiene "
        "su Executor en su hilo; el SwarmService los crea y los mata, nunca los reusa.")
    add_para(doc, "swarm_service.py", bold=True)
    add_para(doc,
        "Es el director de orquesta. Carga fleet.yaml, asegura tablas en BD, crea un "
        "DroneService por dron, inicia telemetria. Mantiene un loop de planificacion "
        "cada PLANNING_PERIOD_S = 5s que lee pedidos pendientes, lee estado de flota, "
        "invoca planner.build_plan(), persiste el plan, lanza Executors y gestiona su "
        "ciclo de vida. Expone API hacia la UI: start, stop, pause_planning, "
        "resume_planning, force_replan, cancel_assignment.")

    add_heading(doc, "3.6 Cambios en la UI", level=2)
    add_bullets(doc, [
        "cliente/main_window.py: sustituir el unico DroneService por un SwarmService; eliminar DRONE_ID hardcoded; suscribirse a senales del swarm; mantener Aceptar como override manual; anadir boton Auto que pausa/reanuda planificacion automatica.",
        "widgets/fleet_panel.py: ya soporta add_drone(id), basta con llamar N veces. Cada DroneCard muestra estado, bateria, pedido actual y ETA.",
        "widgets/map_widget.py: cambiar update_drone_position para aceptar drone_id como primer parametro; mantener un trail por dron con color distinto.",
    ])

    add_heading(doc, "3.7 Concurrencia y robustez", level=2)
    add_para(doc,
        "Cada DroneService vive en su propio thread daemon. SwarmService.start() arranca "
        "un QTimer cada PLANNING_PERIOD_S. El Planner es sincrono y rapido (<= 5s timeout); "
        "se invoca dentro de un QThreadPool para no congelar la UI. Los Executors emiten "
        "senales Qt; la MainWindow solo actualiza la UI cuando llegan al hilo principal por "
        "la cola de eventos de Qt.")
    add_para(doc, "Tabla de fallos y respuestas:", bold=True)
    add_table_simple(doc,
        headers=["Fallo", "Estrategia"],
        rows=[
            ("Perdida de conexion MAVLink", "DroneService marca estado offline; planner lo excluye del plan."),
            ("Energia estimada > umbral mid-flight", "Executor aborta y manda RTL; replanifica al volver."),
            ("Pedido sin destino valido", "Queda unassigned; alerta en UI."),
            ("Conflicto en HUB no resuelto", "Airspace fuerza takeoff_offset += 5s; log de aviso."),
            ("Solver infeasible", "Fallback Hungarian; los huerfanos esperan al siguiente ciclo."),
            ("Crash de un Executor", "SwarmService lo captura; assignment marcado aborted."),
        ], col_widths=[6, 11])

    doc.add_page_break()

    # ---------------- ROADMAP ----------------
    add_heading(doc, "4. Roadmap por fases", level=1)
    add_para(doc,
        "El plan se divide en 6 fases. Cada fase deja el sistema arrancable y "
        "demostrable; ninguna fase deja el repositorio a medias. Si en una fase "
        "concreta se acaba el tiempo, la entrega anterior sigue siendo funcional.")
    add_table_simple(doc,
        headers=["Fase", "Tema", "Estim.", "Bloquea a"],
        rows=[
            ("0", "Pre-flight: tooling, deps, branch", "0.5 d", "Todo"),
            ("1", "Modelo de dominio + persistencia", "1 d", "2, 3, 5"),
            ("2", "Modelo de energia + Planner CVRP", "2 d", "3"),
            ("3", "Executor + Airspace Manager", "2 d", "4, 5"),
            ("4", "SwarmService + integracion SITL multi-dron", "2 d", "5"),
            ("5", "UI multi-dron + portal", "1.5 d", "6"),
            ("6", "Validacion, demo, memoria, video", "1 d", "-"),
        ], col_widths=[1.2, 8, 2, 4])
    add_para(doc, "Total estimado: aproximadamente 10 dias de trabajo dedicado.", italic=True)

    add_heading(doc, "4.1 Fase 0 - Pre-flight", level=2)
    add_para(doc, "Meta: dejar el entorno listo y un branch limpio.")
    add_numbered(doc, [
        "git init && git commit -m 'snapshot pre-swarm' del estado actual.",
        "Crear branch feature/swarm-manager.",
        "Actualizar requirements.txt con ortools>=9.10, numpy>=1.26, pyyaml>=6.0, scipy>=1.11, pytest>=8.0.",
        "pip install -r requirements.txt y verificar que from ortools.constraint_solver import pywrapcp importa sin error en Windows.",
        "Crear tests/ con test_smoke.py y configurar pytest.",
    ])
    add_para(doc, "Criterio de aceptacion: python -m pytest tests/ pasa.", bold=True)

    add_heading(doc, "4.2 Fase 1 - Modelo de dominio + persistencia", level=2)
    add_para(doc, "Meta: la BD conoce drones y assignments; el codigo tiene los DTOs.")
    add_numbered(doc, [
        "Crear modelos/drone.py con DroneSpec, DroneRuntime y Drone.",
        "Crear modelos/assignment.py y modelos/flight_plan.py.",
        "Ampliar negocio/db_manager.py con tablas drones, drone_states, flight_plans, assignments y migracion idempotente.",
        "Crear negocio/fleet_repository.py con list_drones, upsert_drone, upsert_state, save_plan, update_assignment_status.",
        "Crear config/fleet.yaml con 3 drones de specs definidas.",
        "Crear negocio/fleet_bootstrap.py que carga fleet.yaml y rellena la tabla la primera vez.",
        "Tests test_fleet_repository.py.",
    ])
    add_para(doc, "Criterio de aceptacion: tras arrancar la app, la tabla drones tiene 3 filas leidas desde YAML; tests pasan; el mono-dron sigue funcionando.", bold=True)

    add_heading(doc, "4.3 Fase 2 - Energia + Planner", level=2)
    add_para(doc, "Meta: dado un estado de flota y unos pedidos, generar un FlightPlan sin tocar todavia MAVLink.")
    add_numbered(doc, [
        "Crear servicios/energy_model.py con estimate_energy_wh() y estimate_duration_s() puras.",
        "Test test_energy_model.py con 4-5 casos calibrados a mano.",
        "Crear servicios/planner_service.py con build_plan(drones, orders, time_budget_s).",
        "Implementar matriz coste/energia/capacidad y llamada a OR-Tools.",
        "Implementar fallback Hungarian si solucion vacia.",
        "Tests test_planner_service.py con 4 casos sinteticos.",
        "Script manual scripts/run_planner_smoke.py.",
    ])
    add_para(doc, "Criterio de aceptacion: tests pasan en <2s; en script manual con datos reales el plan es coherente.", bold=True)

    add_heading(doc, "4.4 Fase 3 - Executor + AirspaceManager", level=2)
    add_para(doc, "Meta: ejecutar un Assignment extraido de un FlightPlan contra SITL.")
    add_numbered(doc, [
        "Crear servicios/airspace_manager.py con reserve() y release().",
        "Tests test_airspace_manager.py: slots unicos y takeoff_offset_s creciente.",
        "Refactor servicios/drone_service.py para aceptar drone_id en __init__ y emitir senales con drone_id.",
        "Crear servicios/executor_service.py con Executor(QObject) que reproduce el flujo actual parametrizado.",
        "Marcar deprecated DroneService.start_order_delivery.",
        "Script scripts/demo_phase3.py.",
    ])
    add_para(doc, "Criterio de aceptacion: un pedido se entrega end-to-end pasando por todas las capas nuevas.", bold=True)

    add_heading(doc, "4.5 Fase 4 - SwarmService + multi-SITL", level=2)
    add_para(doc, "Meta: 3 drones simulados en paralelo entregando varios pedidos.")
    add_numbered(doc, [
        "Crear servicios/swarm_service.py con bootstrap, dict[str, DroneService], QTimer y lifecycle de Executors.",
        "Cablear SwarmService desde main_window.py; eliminar DRONE_ID hardcoded.",
        "Cablear senales state_changed, progress, order_status_changed.",
        "Documentar arranque multi-SITL en docs/swarm/RUN_SITL.md.",
        "Pruebas con 2 drones y luego con 3.",
        "Manejar caida de un dron mid-flight.",
    ])
    add_para(doc, "Criterio de aceptacion: SITL x3, 5 pedidos por el portal, 3 drones despegan en distintos slots y los pedidos quedan entregado.", bold=True)

    add_heading(doc, "4.6 Fase 5 - UI multi-dron", level=2)
    add_numbered(doc, [
        "DroneCard: barra de bateria y label ETA.",
        "MapWidget: soporte N marcadores y N trails con color por dron.",
        "FleetPanel: actualizacion de _count_lbl y seleccion multiple.",
        "Toolbar: botones Auto/Manual y Re-plan ahora.",
        "Portal cliente: estado del pedido en tiempo real con polling 2s.",
        "Leyenda de colores de drones en el mapa.",
    ])

    add_heading(doc, "4.7 Fase 6 - Validacion + memoria + video", level=2)
    add_numbered(doc, [
        "Definir escenario de validacion reproducible (3 drones, 6 pedidos) en docs/swarm/SCENARIO_DEMO.md.",
        "Instrumentar logging: swarm_service escribe logs/swarm.log.",
        "Ejecutar escenario y capturar pantallas y logs.",
        "Tabla de metricas: makespan, utilizacion, error de estimacion energetica, tiempo de computo del planner.",
        "Grabar video demo 90-120s con voz en off.",
        "Pasar bibliografia al capitulo de marco teorico del TFG.",
        "git tag v2.0-swarm.",
    ])

    add_heading(doc, "4.8 Riesgos y mitigaciones", level=2)
    add_table_simple(doc,
        headers=["Riesgo", "Mitigacion"],
        rows=[
            ("OR-Tools no instala en Windows", "Plan B: solo Hungarian. Documentado como degradacion aceptable."),
            ("SITL multi-instancia no arranca en local", "Usar maquina virtual / WSL2. Comandos en RUN_SITL.md."),
            ("dronLink con bugs en varias conexiones", "Cada Dron() en su thread; si peta, fallback a pymavlink directo."),
            ("Sin tiempo en Fase 4", "Defendible con 2 drones (demostracion de escalabilidad)."),
            ("Modelo de energia inexacto", "Reportar error y citar Zhang 2021."),
        ], col_widths=[7, 10])

    add_heading(doc, "4.9 Definicion de Done para el TFG", level=2)
    add_bullets(doc, [
        "Existe un branch feature/swarm-manager mergeable a main.",
        "pytest pasa todos los tests unitarios.",
        "La aplicacion arranca con python main.py con la flota cargada de YAML.",
        "Demo SITL muestra >=2 drones entregando >=3 pedidos en paralelo.",
        "Memoria del TFG incluye el marco teorico con sus citas.",
        "Existe un video demo.",
    ])

    doc.add_page_break()

    # ---------------- TAREAS ----------------
    add_heading(doc, "5. Lista granular de tareas", level=1)
    add_para(doc,
        "Cada tarea es atomica, ejecutable en una sesion. Esta lista es input directo "
        "para TaskCreate o para una pizarra Kanban. Convencion: (F) = ficheros tocados, "
        "(V) = criterio de verificacion.")

    add_heading(doc, "Fase 0", level=2)
    add_bullets(doc, [
        "T0.1 Commit de snapshot inicial.",
        "T0.2 Crear branch feature/swarm-manager.",
        "T0.3 Actualizar requirements.txt (ortools, numpy, pyyaml, scipy, pytest).",
        "T0.4 Verificar instalacion OR-Tools.",
        "T0.5 Crear tests/__init__.py y tests/test_smoke.py.",
    ])

    add_heading(doc, "Fase 1", level=2)
    add_bullets(doc, [
        "T1.1 Crear modelos/drone.py con DroneSpec, DroneRuntime, Drone.",
        "T1.2 Crear modelos/assignment.py con Assignment.",
        "T1.3 Crear modelos/flight_plan.py con FlightPlan.",
        "T1.4 Ampliar negocio/db_manager.py con tablas drones, drone_states, flight_plans, assignments y migracion idempotente.",
        "T1.5 Anadir estado 'asignado' al CHECK de orders.status.",
        "T1.6 Crear negocio/fleet_repository.py.",
        "T1.7 Crear config/fleet.yaml con 3 drones.",
        "T1.8 Crear negocio/fleet_bootstrap.py.",
        "T1.9 Llamar al bootstrap desde MainWindow.__init__.",
        "T1.10 Tests tests/test_fleet_repository.py.",
    ])

    add_heading(doc, "Fase 2", level=2)
    add_bullets(doc, [
        "T2.1 Crear servicios/energy_model.py.",
        "T2.2 Tests tests/test_energy_model.py.",
        "T2.3 Esqueleto servicios/planner_service.py.",
        "T2.4 Implementar modelo OR-Tools con dimensiones capacity y energy.",
        "T2.5 Mapeo solucion -> list[Assignment].",
        "T2.6 Fallback Hungarian con scipy.optimize.linear_sum_assignment.",
        "T2.7 Tests test_planner_service.py con 4 casos.",
        "T2.8 Script scripts/run_planner_smoke.py.",
    ])

    add_heading(doc, "Fase 3", level=2)
    add_bullets(doc, [
        "T3.1 Crear servicios/airspace_manager.py.",
        "T3.2 Tests tests/test_airspace_manager.py.",
        "T3.3 Refactor DroneService para aceptar drone_id.",
        "T3.4 Crear Executor(QObject) en servicios/executor_service.py.",
        "T3.5 Deprecar DroneService.start_order_delivery.",
        "T3.6 Script scripts/demo_phase3.py.",
    ])

    add_heading(doc, "Fase 4", level=2)
    add_bullets(doc, [
        "T4.1 Crear servicios/swarm_service.py completo.",
        "T4.2 Reemplazar DroneService por SwarmService en main_window.py.",
        "T4.3 Cablear senales SwarmService <-> FleetPanel.",
        "T4.4 Documentar arranque multi-SITL en RUN_SITL.md.",
        "T4.5 Prueba con 2 drones.",
        "T4.6 Prueba con 3 drones.",
        "T4.7 Manejar caida de un dron mid-flight.",
    ])

    add_heading(doc, "Fase 5", level=2)
    add_bullets(doc, [
        "T5.1 DroneCard con bateria y ETA.",
        "T5.2 MapWidget con drone_id en update_drone_position.",
        "T5.3 FleetPanel con _count_lbl y seleccion multiple.",
        "T5.4 Toolbar con botones Auto/Manual y Re-plan ahora.",
        "T5.5 Portal con polling 2s del estado.",
        "T5.6 Leyenda de colores en el mapa.",
    ])

    add_heading(doc, "Fase 6", level=2)
    add_bullets(doc, [
        "T6.1 Documentar escenario en SCENARIO_DEMO.md.",
        "T6.2 Logging a logs/swarm.log.",
        "T6.3 Capturar pantallas y logs del escenario.",
        "T6.4 Tabla de metricas: makespan, utilizacion, error energetico, tiempo solver.",
        "T6.5 Grabar video demo.",
        "T6.6 Pasar bibliografia al TFG.",
        "T6.7 git tag v2.0-swarm.",
    ])

    add_heading(doc, "Tareas transversales", level=2)
    add_bullets(doc, [
        "TX.1 Actualizar README del repo con seccion Swarm Manager.",
        "TX.2 Limpiar business_manager_BACKUP.py y DesktopLAN_BACKUP.py si el tutor lo aprueba.",
        "TX.3 Anotaciones de tipo (from __future__ import annotations) en modulos nuevos.",
        "TX.4 Pre-commit / black / ruff opcional.",
    ])

    doc.add_page_break()

    # ---------------- REFERENCIAS ----------------
    add_heading(doc, "6. Referencias bibliograficas", level=1)
    add_para(doc,
        "Las referencias siguen orden alfabetico por primer autor. Todas las URL son "
        "publicas; las marcadas con (paywall) requieren acceso institucional al "
        "articulo final, pero las versiones preprint suelen estar accesibles en arXiv "
        "o repositorios institucionales.")

    refs = [
        ("Abeywickrama, H. V., Jayawickrama, B. A., He, Y., Dutkiewicz, E.",
         "2018",
         "Battery-Aware Energy Model of Drone Delivery Tasks",
         "ResearchGate preprint",
         "https://www.researchgate.net/publication/330486353_Battery-Aware_Energy_Model_of_Drone_Delivery_Tasks"),
        ("ArduPilot",
         "2024",
         "Multi-Vehicle Flying - Copter",
         "Documentacion oficial ArduPilot",
         "https://ardupilot.org/copter/docs/common-multi-vehicle-flying.html"),
        ("ArduPilot",
         "2024",
         "Mission Planner Swarming (beta)",
         "Documentacion oficial Mission Planner",
         "https://ardupilot.org/planner/docs/swarming.html"),
        ("Bryant, S., Yetkin, H., Hennig, T.",
         "2020",
         "Energy-Constrained Delivery of Goods with Drones Under Varying Wind Conditions",
         "arXiv:2012.08602",
         "https://arxiv.org/pdf/2012.08602"),
        ("Choi, H.-L., Brunet, L., How, J. P.",
         "2009",
         "Consensus-Based Decentralized Auctions for Robust Task Allocation",
         "IEEE T-RO 25(4)",
         "https://www.researchgate.net/publication/228529155_Consensus-Based_Auction_Approaches_for_Decentralized_Task_Assignment"),
        ("Dorling, K., Heinrichs, J., Messier, G., Magierowski, S.",
         "2017",
         "Vehicle Routing Problems for Drone Delivery",
         "IEEE T-SMC 47(1)",
         "https://arxiv.org/pdf/1608.02305"),
        ("Google",
         "2024",
         "OR-Tools - Routing",
         "Documentacion oficial",
         "https://developers.google.com/optimization/routing"),
        ("Lakhdari, A. et al.",
         "2020",
         "Swarm-based Drone-as-a-Service (SDaaS) for Delivery",
         "arXiv:2005.06952",
         "https://arxiv.org/pdf/2005.06952"),
        ("Lakhdari, A. et al.",
         "2022",
         "In-Flight Energy-Driven Composition of Drone Swarm Services",
         "arXiv:2210.17294",
         "https://arxiv.org/pdf/2210.17294"),
        ("Lin, Z. et al.",
         "2023",
         "Optimal delivery route planning for a fleet of heterogeneous drones: A rescheduling-based GA approach",
         "Computers & Industrial Engineering (paywall)",
         "https://www.sciencedirect.com/science/article/abs/pii/S0360835223002036"),
        ("Liu, X., Hao, S., Zhou, Z.",
         "2023",
         "Drone Routing for Drone-Based Delivery Systems: A Review of Trajectory Planning, Charging, and Security",
         "Sensors 23(3):1463",
         "https://www.mdpi.com/1424-8220/23/3/1463"),
        ("MAVLink",
         "2024",
         "MAVLink System and Component ID Assignment",
         "Documentacion oficial",
         "https://mavlink.io/en/services/mavlink_id_assignment.html"),
        ("MAVLink",
         "2024",
         "Mission Protocol",
         "Documentacion oficial",
         "https://mavlink.io/en/services/mission.html"),
        ("Aerospace MDPI (ed.)",
         "2023",
         "Airspace Designs and Operations for UAS Traffic Management at Low Altitude",
         "Aerospace 10(9):737",
         "https://www.mdpi.com/2226-4310/10/9/737"),
        ("Optimal Collaborative Transportation",
         "2023",
         "Optimal Collaborative Transportation for Under-Capacitated VRP using Aerial Drone Swarms",
         "arXiv:2310.02726",
         "https://arxiv.org/pdf/2310.02726"),
        ("Two-Level Clustered CBBA (autores varios)",
         "2025",
         "A Two-Level Clustered Consensus-Based Bundle Algorithm for Dynamic Heterogeneous Multi-UAV Multi-Task Allocation",
         "Sensors 25(21):6738",
         "https://www.mdpi.com/1424-8220/25/21/6738"),
        ("Zhang, J., Campbell, J., Sweeney, D., Hupman, A.",
         "2021",
         "Energy Consumption Models for Delivery Drones: A Comparison and Assessment",
         "Transp. Res. D 90:102668 (paywall) - preprint ResearchGate",
         "https://www.researchgate.net/publication/341945067_Energy_Consumption_Models_for_Delivery_Drones_A_Comparison_and_Assessment"),
        ("Liu, Y. (Wayne State Univ.)",
         "2020",
         "Strategic Deconfliction of Unmanned Aircraft",
         "AIAA / Wayne State preprint",
         "https://yliu.eng.wayne.edu/research/utm_flight_aiaa_final.pdf"),
        ("DroneKit",
         "2024",
         "Setting up a Simulated Vehicle (SITL)",
         "Documentacion oficial DroneKit-Python",
         "https://dronekit-python.readthedocs.io/en/latest/develop/sitl_setup.html"),
        ("A-POMO (AIAA)",
         "2024",
         "Adaptive Policy Optimization for Battery-Constrained Drone Delivery Routing",
         "Journal of Aerospace Information Systems (paywall)",
         "https://arc.aiaa.org/doi/10.2514/1.I011709"),
    ]
    for i, r in enumerate(refs, start=1):
        add_reference(doc, i, r[0], r[1], r[2], r[3], r[4])

    # ---------------- ANEXO ----------------
    doc.add_page_break()
    add_heading(doc, "Anexo A - Comandos de arranque SITL multi-instancia", level=1)
    add_para(doc,
        "Como referencia operativa para la Fase 4 del roadmap, los siguientes comandos "
        "arrancan tres instancias SITL de ArduCopter con SYSID unicos y puertos TCP "
        "diferentes. En Linux/WSL2:")
    add_code(doc,
"""# Terminal 1 - Dron 1 (SYSID=1, TCP:5760)
cd ~/ardupilot/ArduCopter && sim_vehicle.py -v ArduCopter --instance 0 \\
    --sysid 1 --console --map -L Castelldefels

# Terminal 2 - Dron 2 (SYSID=2, TCP:5763)
sim_vehicle.py -v ArduCopter --instance 1 \\
    --sysid 2 --console -L Castelldefels

# Terminal 3 - Dron 3 (SYSID=3, TCP:5766)
sim_vehicle.py -v ArduCopter --instance 2 \\
    --sysid 3 --console -L Castelldefels
""")
    add_para(doc,
        "El fichero config/fleet.yaml enlaza esos puertos a Dron-1/2/3 mediante el "
        "campo connection: 'tcp:127.0.0.1:5760', '...:5763', '...:5766'.")

    add_heading(doc, "Anexo B - Especificacion declarativa de la flota (fleet.yaml)", level=1)
    add_code(doc,
"""drones:
  - id: Dron-1
    sysid: 1
    connection: \"tcp:127.0.0.1:5760\"
    max_payload_kg: 2.0
    battery_capacity_wh: 222     # ~6S 5000mAh
    cruise_speed_mps: 7.0
    base_consumption_w: 180
    payload_consumption_w_per_kg: 22
    home_parking_name: \"Central\"
  - id: Dron-2
    sysid: 2
    connection: \"tcp:127.0.0.1:5763\"
    max_payload_kg: 1.0          # categoria pequena
    battery_capacity_wh: 148
    cruise_speed_mps: 8.0
    base_consumption_w: 140
    payload_consumption_w_per_kg: 30
    home_parking_name: \"Central\"
  - id: Dron-3
    sysid: 3
    connection: \"tcp:127.0.0.1:5766\"
    max_payload_kg: 4.0          # categoria grande
    battery_capacity_wh: 360
    cruise_speed_mps: 6.0
    base_consumption_w: 280
    payload_consumption_w_per_kg: 18
    home_parking_name: \"Central\"
""")

    doc.save(OUTPUT)
    print("OK ->", OUTPUT)


if __name__ == "__main__":
    build()
